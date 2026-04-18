"""Table builders."""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document


def export_table(df: pd.DataFrame, stem: Path) -> None:
    """Export table in CSV, LaTeX, and DOCX formats."""
    stem.parent.mkdir(parents=True, exist_ok=True)
    df.to_csv(stem.with_suffix(".csv"), index=False)
    stem.with_suffix(".tex").write_text(df.to_latex(index=False, float_format="%.3f"), encoding="utf-8")
    doc = Document()
    doc.add_heading(stem.stem, level=1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            cells[i].text = "" if pd.isna(row[col]) else str(row[col])
    doc.save(stem.with_suffix(".docx"))
